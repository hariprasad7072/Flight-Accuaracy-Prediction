"""
Generate the AeroPredict Mini Project Report in DOCX format.
Uses the template structure from MP2911-Mini Project E-Report Template.docx
and content from the AeroPredict conference paper.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_formatted_paragraph(doc, text, font_name='Times New Roman', font_size=12,
                            bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            space_before=0, space_after=6, line_spacing=1.15,
                            first_line_indent=None, caps=False):
    """Add a formatted paragraph."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)

    run = p.add_run(text.upper() if caps else text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    return p

def add_chapter_heading(doc, number, title):
    """Add a chapter heading: CHAPTER N, then the title."""
    add_formatted_paragraph(doc, f'CHAPTER {number}', font_size=14, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=24, space_after=6)
    add_formatted_paragraph(doc, title.upper(), font_size=13, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=12)

def add_section_heading(doc, number, title):
    """Add a section heading like 2.1 Title."""
    add_formatted_paragraph(doc, f'{number} {title}', font_size=13, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6)

def add_body_text(doc, text):
    """Add justified body text."""
    return add_formatted_paragraph(doc, text, font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                   space_after=6, first_line_indent=1.27)

def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.font.bold = True
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    doc.add_paragraph()  # spacing
    return table

def add_image_if_exists(doc, filename, width=5.5, caption=None):
    """Add an image with optional caption."""
    path = os.path.join(BASE_DIR, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width))
        if caption:
            cap_p = doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap_p.add_run(caption)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.italic = True
    else:
        add_formatted_paragraph(doc, f'[Image: {filename} not found]', font_size=10,
                                italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

def add_page_break(doc):
    doc.add_page_break()

def generate_report():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(2.54)

    # ============================
    # TITLE PAGE
    # ============================
    for _ in range(3):
        doc.add_paragraph()

    add_formatted_paragraph(doc, 'AEROPREDICT: AN EXPLAINABLE MACHINE LEARNING FRAMEWORK FOR FLIGHT DELAY PREDICTION USING XGBOOST AND SHAP ANALYSIS',
                            font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    add_formatted_paragraph(doc, 'MP2911 – MINI PROJECT', font_size=13, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    add_formatted_paragraph(doc, 'Submitted by', font_size=12,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    add_formatted_paragraph(doc, 'I Hari Prasad', font_size=12, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_formatted_paragraph(doc, '(Register Number)', font_size=12,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    add_formatted_paragraph(doc, 'in partial fulfillment for the award of the degree of', font_size=12,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    add_formatted_paragraph(doc, 'BACHELOR OF TECHNOLOGY', font_size=13, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_formatted_paragraph(doc, 'in', font_size=12,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_formatted_paragraph(doc, 'COMPUTER SCIENCE AND ENGINEERING', font_size=13, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_formatted_paragraph(doc, '(AI/AIML)', font_size=12, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    add_formatted_paragraph(doc, 'Academic Year 2025-2026', font_size=12,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_formatted_paragraph(doc, 'EVEN SEMESTER', font_size=12, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    add_formatted_paragraph(doc, 'DIVISION OF ARTIFICIAL INTELLIGENCE AND MACHINE LEARNING', font_size=11, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_formatted_paragraph(doc, 'SCHOOL OF COMPUTER SCIENCE AND TECHNOLOGY', font_size=11, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    add_formatted_paragraph(doc, 'April 2026', font_size=12, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    add_page_break(doc)

    # ============================
    # BONAFIDE CERTIFICATE
    # ============================
    add_formatted_paragraph(doc, 'BONAFIDE CERTIFICATE', font_size=14, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=24, space_after=18)

    add_body_text(doc, 'This is to certify that the mini project report entitled "AeroPredict: An Explainable Machine Learning Framework for Flight Delay Prediction Using XGBoost and SHAP Analysis" is a bonafide work done during the even semester of the academic year 2025-2026 by')

    add_formatted_paragraph(doc, 'I Hari Prasad (Register Number)', font_size=12, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=6)

    add_body_text(doc, 'in partial fulfillment of the requirements for the award of the degree of Bachelor of Technology in Computer Science and Engineering (AI/AIML) of Karunya Institute of Technology and Sciences.')

    for _ in range(4):
        doc.add_paragraph()

    add_formatted_paragraph(doc, 'Signature of the Supervisor', font_size=12,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)
    add_formatted_paragraph(doc, 'Mrs. Sophia S', font_size=12, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    add_formatted_paragraph(doc, 'Division of Artificial Intelligence and Machine Learning', font_size=12,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=18)

    add_formatted_paragraph(doc, 'Submitted for the Viva Voce held on ___________________', font_size=12,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=12)

    add_page_break(doc)

    # ============================
    # DECLARATION
    # ============================
    add_formatted_paragraph(doc, 'DECLARATION', font_size=14, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=24, space_after=18)

    add_body_text(doc, 'I, the undersigned, hereby declare that the mini project titled "AeroPredict: An Explainable Machine Learning Framework for Flight Delay Prediction Using XGBoost and SHAP Analysis" is an original work carried out by me. I am fully responsible for the planning, analysis, design, development, and implementation of the Mini Project described in this report.')

    add_body_text(doc, 'To the best of my knowledge and belief:')

    declarations = [
        'The mini project requirements have been carefully studied, analyzed, and accurately documented in this report.',
        'The mini project has been developed in accordance with the defined specifications, functionalities, and objectives.',
        'Appropriate methodologies, coding standards, and security practices have been followed to ensure the quality and reliability of the mini project.',
        'Sufficient importance has been given to testing, debugging, and validation to deliver a stable and error-free solution.',
        'Any issues, ambiguities, or changes encountered during the mini project development have been identified and addressed appropriately.',
        'The mini project was successfully completed within the stipulated time frame, achieving the intended outcomes while maintaining the required quality standards.'
    ]

    for decl in declarations:
        p = doc.add_paragraph(style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.15
        # Clear default and add formatted run
        p.clear()
        run = p.add_run(decl)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Signature of student                                        Signature of the Supervisor')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    add_page_break(doc)

    # ============================
    # TABLE OF CONTENTS
    # ============================
    add_formatted_paragraph(doc, 'TABLE OF CONTENTS', font_size=14, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=24, space_after=18)

    toc_items = [
        ('', 'Abstract', ''),
        ('1', 'Introduction', ''),
        ('2', 'Literature Survey', ''),
        ('3', 'Proposed Methodology', ''),
        ('', '3.1 Dataset Description', ''),
        ('', '3.2 Feature Engineering', ''),
        ('', '3.3 Data Preprocessing', ''),
        ('', '3.4 Model Architectures', ''),
        ('', '3.5 SHAP Explainability Framework', ''),
        ('', '3.6 System Architecture', ''),
        ('4', 'Results and Discussions', ''),
        ('', '4.1 Experimental Setup', ''),
        ('', '4.2 Evaluation Metrics', ''),
        ('', '4.3 Model Performance Comparison', ''),
        ('', '4.4 Confusion Matrix Analysis', ''),
        ('', '4.5 Cross-Validation Analysis', ''),
        ('', '4.6 ROC Curves and Precision-Recall Analysis', ''),
        ('', '4.7 Lift Chart Analysis', ''),
        ('5', 'Conclusion and Future Work', ''),
        ('6', 'References', ''),
        ('', 'Appendix A – Sample Code', ''),
        ('', 'Appendix B – Output Screenshots', ''),
        ('', 'Appendix C – Outcome Achieved', ''),
    ]

    toc_table = doc.add_table(rows=len(toc_items), cols=3)
    toc_table.style = 'Table Grid'
    for i, (num, title, page) in enumerate(toc_items):
        for j, val in enumerate([num, title, page]):
            cell = toc_table.rows[i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            if num and j == 1:
                run.font.bold = True
        toc_table.rows[i].cells[0].width = Inches(0.5)
        toc_table.rows[i].cells[1].width = Inches(4.5)
        toc_table.rows[i].cells[2].width = Inches(1.0)

    add_page_break(doc)

    # ============================
    # ABSTRACT
    # ============================
    add_formatted_paragraph(doc, 'ABSTRACT', font_size=14, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=24, space_after=12)

    add_body_text(doc, 'Flight delays pose significant operational and economic challenges to the aviation industry, affecting airlines, airports, and passengers worldwide. Accurate prediction of flight delays can enable proactive decision-making and resource optimization. This project presents AeroPredict, an explainable machine learning framework for binary flight delay classification that integrates high-accuracy predictive modeling with post-hoc interpretability through SHapley Additive exPlanations (SHAP).')

    add_body_text(doc, 'A comprehensive pipeline is developed encompassing feature engineering with 25 multi-dimensional input features — spanning meteorological conditions (temperature, precipitation, wind speed, visibility, humidity, pressure), temporal patterns (month, day of week, departure hour), route characteristics (distance, airport congestion index), and operational parameters (aircraft age, passenger load factor, fuel weight, cargo weight, ground crew efficiency) — followed by feature selection and standardized scaling.')

    add_body_text(doc, 'Six distinct models are trained and rigorously evaluated on a dataset of 100,000+ flight records: XGBoost, LightGBM, Random Forest, Deep Neural Network (with Batch Normalization and Dropout), Soft Voting Ensemble, and Stacking Ensemble with Logistic Regression meta-learner. Experimental results demonstrate that all models surpass the 90% accuracy threshold, with the Neural Network achieving the highest accuracy of 98.51% (AUC-ROC: 0.9883) and XGBoost attaining 95.56% accuracy (AUC-ROC: 0.9842).')

    add_body_text(doc, 'Five-fold stratified cross-validation confirms model generalizability with minimal variance. To address the black-box nature of ensemble methods, SHAP-based feature attribution is integrated to provide transparent, per-prediction explanations identifying the top contributing factors to each delay prediction. The framework is deployed as an interactive web application (AeroPredict AI) built on Flask, enabling real-time delay prediction with dynamic weather integration, automatic route parameter computation via Haversine distance calculation, and visual SHAP-driven explanations.')

    add_body_text(doc, 'Comprehensive evaluation through confusion matrices, ROC curves, precision-recall curves, calibration curves, learning curves, and radar charts validates the robustness and reliability of the proposed system, establishing AeroPredict as a practical, interpretable tool for aviation delay risk assessment.')

    add_formatted_paragraph(doc, 'Keywords: Flight Delay Prediction, XGBoost, SHAP, Explainable AI (XAI), Ensemble Learning, Deep Neural Network, Aviation Analytics, Machine Learning',
                            font_size=11, bold=True, italic=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=12)

    add_page_break(doc)

    # ============================
    # CHAPTER 1: INTRODUCTION
    # ============================
    add_chapter_heading(doc, '1', 'INTRODUCTION')

    add_body_text(doc, 'Air transportation serves as a vital pillar of global connectivity, facilitating the movement of over 4.5 billion passengers annually and contributing approximately 3.5% to the world\'s gross domestic product. Despite its significance, the aviation industry continues to face the persistent challenge of flight delays, which propagate across interconnected airline networks and result in economic losses exceeding $30 billion annually in the United States alone. These delays degrade passenger satisfaction, escalate operational costs for airlines, and impose considerable strain on airport infrastructure and air traffic management systems. Consequently, the accurate and timely prediction of flight delays has become a critical priority for all stakeholders within the aviation ecosystem.')

    add_body_text(doc, 'Flight delays are influenced by a complex interplay of heterogeneous factors, including adverse weather conditions, air traffic congestion, mechanical failures, crew scheduling limitations, and regulatory constraints. The multifactorial and inherently stochastic nature of these variables makes delay prediction a non-trivial problem, largely unsuitable for conventional rule-based or statistical approaches. In recent years, machine learning (ML) techniques have shown remarkable potential in modeling the complex, non-linear relationships embedded in aviation operational data. Supervised learning methods — including decision trees, support vector machines, and gradient boosting frameworks — have been increasingly applied to flight delay classification, though with varying levels of effectiveness.')

    add_body_text(doc, 'Among ensemble learning methods, gradient boosting algorithms such as XGBoost and LightGBM have consistently demonstrated state-of-the-art performance across diverse predictive tasks. Their capacity to handle heterogeneous feature spaces, capture intricate feature interactions, and mitigate overfitting through built-in regularization makes them particularly effective for flight delay classification. Similarly, deep neural networks (DNNs) offer the ability to learn hierarchical feature representations from large-scale datasets. However, the superior predictive accuracy of these models often comes at the expense of interpretability — a critical limitation in safety-critical domains such as aviation, where transparent and accountable decision-making is essential.')

    add_body_text(doc, 'Explainable Artificial Intelligence (XAI) has emerged as a promising paradigm to bridge this gap. SHapley Additive exPlanations (SHAP), grounded in cooperative game theory, provides a mathematically rigorous framework for quantifying the contribution of each input feature to an individual prediction. By computing Shapley values, SHAP delivers both global feature importance rankings and local, instance-level explanations, enabling stakeholders to understand precisely why a given flight is predicted to experience a delay.')

    add_section_heading(doc, '1.1', 'Problem Statement')

    add_body_text(doc, 'Despite the expanding body of research on flight delay prediction, several notable gaps remain unaddressed. First, many studies evaluate only a narrow range of ML models, lacking systematic comparison under uniform experimental settings. Second, feature engineering is often limited to temporal and carrier-level attributes, while operationally significant factors — such as aircraft age, airport congestion, passenger load, ground crew efficiency, and cargo weight — remain underexplored. Third, the integration of post-hoc explainability methods into a fully deployable, end-to-end prediction system is rarely realized. Fourth, the transition from a research prototype to an accessible, interactive tool for non-technical users receives insufficient attention.')

    add_section_heading(doc, '1.2', 'Objectives')

    add_body_text(doc, 'To address these challenges, this project presents AeroPredict, a comprehensive and explainable ML framework for flight delay prediction. The key objectives of this work are:')

    objectives = [
        'To conduct a rigorous comparative evaluation of six ML models — XGBoost, LightGBM, Random Forest, Deep Neural Network, Soft Voting Ensemble, and Stacking Ensemble — on a dataset of over 100,000 flight records.',
        'To engineer a rich 25-dimensional feature space incorporating meteorological, temporal, route-based, and operational variables to capture the full spectrum of delay-inducing factors.',
        'To integrate SHAP-based explainability to deliver per-prediction feature attribution, transforming the classifier from a black-box model into a transparent decision-support tool.',
        'To deploy the complete framework as AeroPredict AI, an interactive Flask-based web application featuring real-time weather integration, automatic route distance computation, and visual SHAP-driven explanations.',
        'To validate the framework through a comprehensive evaluation suite including confusion matrices, ROC curves, precision-recall curves, calibration plots, learning curves, and 5-fold stratified cross-validation.'
    ]

    for i, obj in enumerate(objectives, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(1.27)
        run = p.add_run(f'{i}. {obj}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    add_page_break(doc)

    # ============================
    # CHAPTER 2: LITERATURE SURVEY
    # ============================
    add_chapter_heading(doc, '2', 'LITERATURE SURVEY')

    add_body_text(doc, 'The prediction of flight delays has attracted significant research attention over the past two decades, driven by the growing availability of aviation data and advances in computational intelligence. This chapter reviews the existing body of work across three thematic areas: traditional and machine learning approaches to flight delay prediction, ensemble and gradient boosting methods, and explainable artificial intelligence in aviation.')

    add_section_heading(doc, '2.1', 'Flight Delay Prediction Using Machine Learning')

    add_body_text(doc, 'Early approaches to flight delay prediction relied on statistical methods such as linear regression, logistic regression, and time-series analysis. Rebollo and Balakrishnan [1] developed a predictive model using Random Forest and regression techniques to characterize air traffic delays in the National Airspace System, achieving moderate accuracy by incorporating temporal and spatial propagation features. Mueller and Chatterji proposed a ground delay program planning model based on weather capacity forecasts, demonstrating the critical role of meteorological factors in delay prediction.')

    add_body_text(doc, 'With the advent of machine learning, researchers began exploring more sophisticated classification algorithms. Choi et al. [3] applied gradient boosting and random forest classifiers to predict weather-induced flight delays, incorporating features such as wind speed, visibility, and precipitation. Their results highlighted the superiority of ensemble methods over single-model approaches. Manna et al. [4] employed support vector machines (SVM) and k-nearest neighbours (KNN) for delay classification using the Bureau of Transportation Statistics (BTS) dataset, reporting accuracies in the range of 80–85%. Chakrabarty [5] applied a Naive Bayes classifier combined with decision tree-based feature selection to predict departure delays.')

    add_body_text(doc, 'More recently, Gui et al. [6] proposed a flight delay prediction model using a combination of deep neural networks and recurrent architectures, leveraging sequential temporal patterns in flight operations. Kim et al. [7] developed a deep learning framework incorporating convolutional layers for spatial feature extraction from airport network data, achieving an accuracy of approximately 87%. Hao et al. [8] utilized long short-term memory (LSTM) networks to capture delay propagation across connecting flights. While these studies advanced prediction accuracy, they predominantly focused on narrow feature sets and lacked systematic multi-model comparison.')

    add_section_heading(doc, '2.2', 'Ensemble Learning and Gradient Boosting Methods')

    add_body_text(doc, 'Ensemble learning methods have emerged as dominant approaches in structured data classification tasks. The foundational work of Breiman [9] on Random Forests demonstrated that aggregating predictions from multiple decorrelated decision trees yields robust and generalizable classifiers. Subsequently, gradient boosting frameworks extended this paradigm by sequentially constructing additive models that optimize a differentiable loss function.')

    add_body_text(doc, 'Chen and Guestrin [10] introduced XGBoost (Extreme Gradient Boosting), which incorporates regularization terms, column subsampling, and efficient split-finding algorithms to deliver state-of-the-art performance. Ke et al. [12] proposed LightGBM, which employs gradient-based one-side sampling (GOSS) and exclusive feature bundling (EFB) to achieve substantially faster training times while maintaining competitive accuracy.')

    add_body_text(doc, 'In the context of flight delay prediction, Ye et al. [13] compared XGBoost, LightGBM, and CatBoost on the U.S. domestic flight dataset, reporting that XGBoost achieved 92.4% accuracy. Jiang and Zhang [14] employed a stacking ensemble combining gradient boosting with logistic regression as the meta-learner, demonstrating that ensemble strategies consistently outperform individual base models by 2–4% in accuracy.')

    add_section_heading(doc, '2.3', 'Explainable Artificial Intelligence in Aviation')

    add_body_text(doc, 'The deployment of machine learning models in safety-critical domains necessitates transparency in model reasoning. Lundberg and Lee [16] introduced SHAP (SHapley Additive exPlanations), a unified framework grounded in Shapley values from cooperative game theory. SHAP provides consistent, locally faithful feature attributions for any predictive model, enabling both global importance ranking and instance-level explanations.')

    add_body_text(doc, 'In aviation applications, Fernandes et al. [18] applied SHAP to interpret gradient boosting models predicting en-route flight trajectory deviations. Wang et al. [19] utilized SHAP values to explain ensemble delay predictions at major Chinese airports, identifying airport congestion and departure time as consistently dominant factors. Despite these advances, the integration of SHAP-based explainability into a fully deployable, interactive prediction system remains largely unexplored in the literature.')

    add_section_heading(doc, '2.4', 'Comparison with Existing Literature')

    # Literature comparison table
    lit_headers = ['Study', 'Models Used', 'Features', 'Accuracy', 'XAI', 'Deployed']
    lit_rows = [
        ['Rebollo & Balakrishnan [1]', 'RF, Regression', 'Temporal, spatial', '~82%', '✗', '✗'],
        ['Choi et al. [3]', 'GB, RF', 'Weather, temporal', '~85%', '✗', '✗'],
        ['Kim et al. [7]', 'CNN, DNN', 'Airport network', '~87%', '✗', '✗'],
        ['Ye et al. [13]', 'XGB, LGBM, CatBoost', 'Temporal, carrier', '~92%', '✗', '✗'],
        ['Jiang & Zhang [14]', 'Stacking Ensemble', 'Temporal, route', '~91%', '✗', '✗'],
        ['Wang et al. [19]', 'GB, RF', 'Weather, airport', '~89%', 'SHAP', '✗'],
        ['AeroPredict (Ours)', '6 Models', '25 features', '98.51%', 'SHAP', '✓'],
    ]
    add_table(doc, lit_headers, lit_rows)

    add_body_text(doc, 'As evident from the table above, the proposed work distinguishes itself across five dimensions: (i) the most comprehensive multi-model comparison including both ensemble and deep learning methods; (ii) the richest feature set incorporating underexplored operational variables; (iii) the highest reported accuracy; (iv) integrated SHAP-based explainability; and (v) deployment as a fully functional, interactive web application.')

    add_page_break(doc)

    # ============================
    # CHAPTER 3: PROPOSED METHODOLOGY
    # ============================
    add_chapter_heading(doc, '3', 'PROPOSED METHODOLOGY')

    add_body_text(doc, 'This chapter presents the end-to-end methodology of the AeroPredict framework, encompassing dataset construction, feature engineering, preprocessing, model architectures, SHAP-based explainability, and system architecture.')

    add_section_heading(doc, '3.1', 'Dataset Description')

    add_body_text(doc, 'The AeroPredict framework is trained and evaluated on a large-scale dataset comprising over 100,000 flight records. Each record represents a single scheduled flight operation and is annotated with a binary target label: delayed (1) or on-time (0). The dataset exhibits a realistic class distribution with approximately 60% on-time flights and 40% delayed flights, reflecting the inherent class imbalance observed in real-world aviation data.')

    add_body_text(doc, 'Each flight record is characterized by 25 input features spanning five distinct categories: temporal attributes, meteorological conditions, route characteristics, aircraft and operational parameters, and encoded categorical identifiers. The dataset is partitioned into training (80%) and testing (20%) subsets using stratified sampling to preserve the class distribution across both splits.')

    add_section_heading(doc, '3.2', 'Feature Engineering')

    add_body_text(doc, 'A comprehensive set of 25 input features is engineered to capture multi-dimensional aspects of flight operations. The complete feature set is described in the table below:')

    # Feature set table
    feat_headers = ['#', 'Feature Name', 'Category', 'Description', 'Range/Type']
    feat_rows = [
        ['1', 'Month', 'Temporal', 'Month of departure', '1–12'],
        ['2', 'Day_of_Week', 'Temporal', 'Day of the week', '1–7'],
        ['3', 'Departure_Hour', 'Temporal', 'Scheduled departure hour', '0–23'],
        ['4', 'Departure_Minute', 'Temporal', 'Scheduled departure minute', '0–59'],
        ['5', 'Scheduled_Arrival_Time', 'Temporal', 'Estimated arrival (min from midnight)', '0–1440'],
        ['6', 'Is_Weekend', 'Temporal', 'Weekend indicator flag', '0 or 1'],
        ['7', 'Is_Holiday', 'Temporal', 'Public holiday indicator flag', '0 or 1'],
        ['8', 'Temperature_C', 'Meteorological', 'Ambient temperature at departure airport', '−10 to 40°C'],
        ['9', 'Precipitation_mm', 'Meteorological', 'Precipitation at departure time', '0–50 mm'],
        ['10', 'Wind_Speed_kmh', 'Meteorological', 'Surface wind speed', '0–80 km/h'],
        ['11', 'Visibility_km', 'Meteorological', 'Horizontal visibility', '1–20 km'],
        ['12', 'Humidity_Pct', 'Meteorological', 'Relative humidity percentage', '20–100%'],
        ['13', 'Pressure_hPa', 'Meteorological', 'Barometric pressure', '980–1050 hPa'],
        ['14', 'Distance_Miles', 'Route', 'Great-circle distance between airports', '100–3000 mi'],
        ['15', 'Airport_Congestion_Index', 'Route', 'Traffic density at departure airport', '0.0–1.0'],
        ['16', 'Num_Connections', 'Route', 'Number of connecting segments', '0–2'],
        ['17', 'Aircraft_Age_Years', 'Operational', 'Age of the aircraft', '0–30'],
        ['18', 'Passenger_Load_Factor', 'Operational', 'Percentage of seat occupancy', '0.5–1.0'],
        ['19', 'Fuel_Weight_kg', 'Operational', 'Total fuel weight loaded', '5,000–50,000 kg'],
        ['20', 'Cargo_Weight_kg', 'Operational', 'Total cargo weight carried', '1,000–15,000 kg'],
        ['21', 'Ground_Crew_Efficiency', 'Operational', 'Ground handling performance score', '0.0–1.0'],
        ['22', 'Previous_Delay_Risk', 'Operational', 'Historical delay probability', '0.0–1.0'],
        ['23', 'Airline_Encoded', 'Categorical', 'Integer-encoded airline carrier', '0–14'],
        ['24', 'Origin_Encoded', 'Categorical', 'Integer-encoded origin airport', '0–299'],
    ]
    add_table(doc, feat_headers, feat_rows)

    add_body_text(doc, 'Additionally, interaction features and polynomial features are generated during training, including pairwise multiplication of top-correlated features, squared terms, and row-wise statistical aggregates (mean, standard deviation, max, range), expanding the initial feature space to enhance model expressiveness.')

    add_section_heading(doc, '3.3', 'Data Preprocessing')

    add_body_text(doc, 'The preprocessing pipeline consists of three stages applied sequentially:')

    add_body_text(doc, '1) Missing Value Imputation: A median-based imputation strategy is employed using SimpleImputer to handle any missing values in the dataset. Median imputation is selected over mean imputation due to its robustness to outliers, which are common in meteorological and operational data.')

    add_body_text(doc, '2) Feature Selection: A Random Forest–based feature selection mechanism is applied using SelectFromModel with a threshold set to the mean feature importance. This step removes low-importance and redundant features, reducing noise and improving model generalization.')

    add_body_text(doc, '3) Feature Scaling: StandardScaler is applied to standardize all retained features to zero mean and unit variance. Scaling is critical for distance-sensitive algorithms (Neural Network) and ensures equitable feature contribution across models with heterogeneous value ranges.')

    add_body_text(doc, 'The preprocessed data is then split into training (80%) and test (20%) sets using stratified random sampling with a fixed seed (SEED = 42) to ensure reproducibility and consistent class distribution.')

    add_section_heading(doc, '3.4', 'Model Architectures')

    add_body_text(doc, 'Six diverse models spanning three paradigms — gradient boosting, bagging, deep learning, and meta-ensemble — are trained and evaluated.')

    add_body_text(doc, '3.4.1 XGBoost (Extreme Gradient Boosting): XGBoost constructs an additive ensemble of decision trees by sequentially minimizing a regularized objective function. The model is configured with 600 estimators, a learning rate of 0.03, maximum depth of 7, and L1/L2 regularization (α = 0.05, λ = 1.0). Column and row subsampling (80%) is applied at each boosting iteration to reduce overfitting.')

    add_body_text(doc, '3.4.2 LightGBM (Light Gradient Boosting Machine): LightGBM employs leaf-wise tree growth with Gradient-based One-Side Sampling (GOSS) for faster convergence. It is configured with 600 estimators, learning rate of 0.03, 63 leaves, maximum depth of 8, and balanced class weights.')

    add_body_text(doc, '3.4.3 Random Forest: An optimized Random Forest classifier with 400 trees, maximum depth of 20, sqrt feature subsampling, and balanced class weights is trained. Minimum samples per split and leaf are set to 4 and 2, respectively.')

    add_body_text(doc, '3.4.4 Deep Neural Network (DNN): A fully connected feedforward neural network is constructed with layers of 256, 128, 64, 32, and 1 neurons. Batch Normalization and Dropout (0.3, 0.3, 0.2) are applied for regularization. The network is trained using the Adam optimizer (lr = 0.001) with binary cross-entropy loss, Early Stopping (patience = 5), and ReduceLROnPlateau callbacks. Training runs for up to 50 epochs with a batch size of 128.')

    add_body_text(doc, '3.4.5 Soft Voting Ensemble: A Voting Classifier with soft voting is constructed by combining the trained XGBoost, LightGBM, and Random Forest models. Each base model contributes its predicted class probabilities, which are averaged to produce the final prediction.')

    add_body_text(doc, '3.4.6 Stacking Ensemble: A Stacking Classifier combines three base learners — XGBoost, LightGBM, and Random Forest — with a Logistic Regression meta-learner (C = 1.0). The stacking is performed with 5-fold cross-validation to generate out-of-fold meta-features.')

    # Hyperparameter table
    hyp_headers = ['Model', 'Key Hyperparameters']
    hyp_rows = [
        ['XGBoost', 'n_estimators=600, lr=0.03, depth=7, subsample=0.8, α=0.05, λ=1.0'],
        ['LightGBM', 'n_estimators=600, lr=0.03, leaves=63, depth=8, α=0.1, λ=1.0'],
        ['Random Forest', 'n_estimators=400, depth=20, min_split=4, min_leaf=2, balanced'],
        ['DNN', 'Layers=[256,128,64,32,1], Dropout=[0.3,0.3,0.2], Adam lr=0.001, epochs=50'],
        ['Voting Ensemble', 'Base: XGB + LGBM + RF, Strategy: Soft voting'],
        ['Stacking Ensemble', 'Base: XGB + LGBM + RF, Meta: LogReg (C=1.0), CV=5'],
    ]
    add_table(doc, hyp_headers, hyp_rows)

    add_section_heading(doc, '3.5', 'SHAP Explainability Framework')

    add_body_text(doc, 'To address the interpretability requirements of aviation decision-making, SHapley Additive exPlanations (SHAP) are integrated into the prediction pipeline. A TreeExplainer is instantiated from the trained XGBoost model, enabling exact and computationally efficient Shapley value computation for tree-based ensembles.')

    add_body_text(doc, 'For each prediction instance, SHAP produces a vector of attribution values φ = {φ₁, φ₂, ..., φ₂₅}, where each φᵢ quantifies the marginal contribution of feature i to the deviation of the prediction from the base rate. Positive Shapley values indicate features that push the prediction toward "delayed," while negative values indicate features that push toward "on-time."')

    add_body_text(doc, 'The top-5 contributing features (ranked by |φᵢ|) are extracted for each prediction and presented with the feature name, feature value, impact direction (increased or decreased delay probability), and importance magnitude.')

    add_section_heading(doc, '3.6', 'System Architecture')

    add_body_text(doc, 'The AeroPredict system operates in two phases:')

    add_body_text(doc, 'Offline Phase (Training): The dataset is preprocessed, features are engineered and selected, all six models are trained and evaluated, and the best-performing XGBoost model along with the fitted StandardScaler and SHAP TreeExplainer are serialized and persisted to disk.')

    add_body_text(doc, 'Online Phase (Inference): The Flask web server loads the pre-trained artifacts at startup. When a user submits a prediction request, the auto-fill module computes route distance (Haversine formula), fetches location-aware weather data, estimates operational parameters, and constructs the complete 25-feature vector. The vector is scaled, passed to the XGBoost model for probability estimation, and simultaneously processed by the SHAP explainer. The combined prediction and explanation are returned as a JSON response.')

    # Technology stack table
    tech_headers = ['Component', 'Technology']
    tech_rows = [
        ['Backend', 'Python Flask (REST API)'],
        ['ML Model', 'XGBoost (pre-trained, serialized as JSON)'],
        ['Explainability', 'SHAP TreeExplainer (serialized via Joblib)'],
        ['Preprocessing', 'StandardScaler (pre-fitted, serialized via Joblib)'],
        ['Frontend', 'HTML5, CSS3, JavaScript'],
    ]
    add_table(doc, tech_headers, tech_rows)

    add_page_break(doc)

    # ============================
    # CHAPTER 4: RESULTS AND DISCUSSIONS
    # ============================
    add_chapter_heading(doc, '4', 'RESULTS AND DISCUSSIONS')

    add_body_text(doc, 'This chapter presents the comprehensive evaluation of all six models trained within the AeroPredict framework. The experimental setup, evaluation metrics, and detailed comparative results are discussed across multiple performance dimensions.')

    add_section_heading(doc, '4.1', 'Experimental Setup')

    add_body_text(doc, 'All experiments are conducted on a dataset of 100,000 flight records, split into 80,000 training samples and 20,000 test samples using stratified random sampling (seed = 42). The class distribution comprises approximately 59.8% on-time (class 0, n = 11,965) and 40.2% delayed (class 1, n = 8,035) flights in the test set. All models are trained on identical training data and evaluated on the same held-out test set to ensure a fair and unbiased comparison. The implementation utilizes Python 3.x with scikit-learn, XGBoost, LightGBM, TensorFlow/Keras, and SHAP libraries.')

    add_section_heading(doc, '4.2', 'Evaluation Metrics')

    add_body_text(doc, 'Model performance is assessed using the following metrics:')
    metrics = [
        'Accuracy: Overall proportion of correct predictions.',
        'Precision: Ratio of true delayed predictions to all predicted delays, measuring false alarm rate.',
        'Recall (Sensitivity): Ratio of correctly identified delays to all actual delays, measuring miss rate.',
        'F1-Score: Harmonic mean of precision and recall, balancing both metrics.',
        'AUC-ROC: Area under the Receiver Operating Characteristic curve, measuring discrimination ability across all classification thresholds.'
    ]
    for m in metrics:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Cm(1.27)
        run = p.add_run(f'• {m}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    add_section_heading(doc, '4.3', 'Model Performance Comparison')

    # Performance table
    perf_headers = ['Model', 'Accuracy', 'Precision', 'Recall', 'F1-Score', 'AUC-ROC']
    perf_rows = [
        ['Neural Network', '0.9851', '0.9819', '0.9811', '0.9815', '0.9883'],
        ['XGBoost', '0.9556', '0.9343', '0.9568', '0.9454', '0.9842'],
        ['Voting Ensemble', '0.9544', '0.9362', '0.9512', '0.9436', '0.9838'],
        ['LightGBM', '0.9536', '0.9314', '0.9548', '0.9430', '0.9838'],
        ['Stacking Ensemble', '0.9526', '0.9440', '0.9376', '0.9408', '0.9833'],
        ['Random Forest', '0.9450', '0.9326', '0.9302', '0.9314', '0.9802'],
    ]
    add_table(doc, perf_headers, perf_rows)

    add_body_text(doc, 'All six models surpass the 90% accuracy threshold, confirming the effectiveness of the engineered feature set. The Neural Network achieves the highest accuracy of 98.51% with an AUC-ROC of 0.9883, followed by XGBoost at 95.56%. Notably, the Stacking Ensemble achieves the highest precision (0.9440) among ensemble models, indicating the lowest false alarm rate, while XGBoost and LightGBM exhibit the highest recall (>0.955), minimizing missed delay predictions.')

    add_image_if_exists(doc, 'model_comparison.png', 5.0, 'Figure 4.1: Model Performance Comparison Bar Chart')
    add_image_if_exists(doc, 'model_performance_radar.png', 5.0, 'Figure 4.2: Multi-Metric Performance Radar Chart')

    add_section_heading(doc, '4.4', 'Confusion Matrix Analysis')

    conf_headers = ['Model', 'TN', 'FP', 'FN', 'TP', 'Total Errors']
    conf_rows = [
        ['Neural Network', '11,820', '145', '152', '7,883', '297'],
        ['XGBoost', '11,424', '541', '347', '7,688', '888'],
        ['Voting Ensemble', '11,444', '521', '392', '7,643', '913'],
        ['LightGBM', '11,400', '565', '363', '7,672', '928'],
        ['Stacking Ensemble', '11,518', '447', '501', '7,534', '948'],
        ['Random Forest', '11,425', '540', '561', '7,474', '1,101'],
    ]
    add_table(doc, conf_headers, conf_rows)

    add_body_text(doc, 'The Neural Network achieves the fewest misclassifications (297 out of 20,000), while XGBoost exhibits the lowest false negative count (347) among ensemble models, prioritizing delay detection. The Stacking Ensemble produces the fewest false positives (447), reflecting a conservative decision boundary from its Logistic Regression meta-learner.')

    add_image_if_exists(doc, 'confusion_matrices.png', 5.0, 'Figure 4.3: Confusion Matrices for All Models')

    add_section_heading(doc, '4.5', 'Cross-Validation Analysis')

    cv_headers = ['Model', 'Fold 1', 'Fold 2', 'Fold 3', 'Fold 4', 'Fold 5', 'Mean ± Std']
    cv_rows = [
        ['XGBoost', '0.9561', '0.9548', '0.9572', '0.9539', '0.9565', '0.9557 ± 0.0012'],
        ['LightGBM', '0.9540', '0.9528', '0.9551', '0.9522', '0.9546', '0.9537 ± 0.0011'],
        ['Random Forest', '0.9455', '0.9441', '0.9468', '0.9437', '0.9459', '0.9452 ± 0.0012'],
        ['Voting Ensemble', '0.9548', '0.9535', '0.9559', '0.9530', '0.9552', '0.9545 ± 0.0011'],
        ['Stacking Ensemble', '0.9531', '0.9518', '0.9540', '0.9515', '0.9536', '0.9528 ± 0.0010'],
    ]
    add_table(doc, cv_headers, cv_rows)

    add_body_text(doc, 'All models exhibit standard deviations below ±0.0013, confirming minimal variance and strong generalizability across data partitions. XGBoost achieves the highest mean CV accuracy (0.9557) among ensemble models, with its test set accuracy (0.9556) closely matching the CV mean — a strong indicator that the reported performance will generalize to unseen real-world data.')

    add_image_if_exists(doc, 'cross_validation_analysis.png', 5.0, 'Figure 4.4: Cross-Validation Analysis')

    add_section_heading(doc, '4.6', 'ROC Curves and Precision-Recall Analysis')

    add_body_text(doc, 'The ROC curves plot the true positive rate against the false positive rate across all classification thresholds. All models achieve AUC-ROC values above 0.98, with the Neural Network achieving the highest at 0.9883.')

    add_image_if_exists(doc, 'roc_curves.png', 5.0, 'Figure 4.5: ROC Curves for All Models')

    add_body_text(doc, 'The precision-recall curves plot precision against recall at varying classification thresholds, providing evaluation that is particularly informative for the delayed class (minority class, 40.2% of test set). All models maintain high precision (>0.93) even at elevated recall levels, indicating that increasing delay detection does not cause excessive false alarms.')

    add_image_if_exists(doc, 'precision_recall_curves.png', 5.0, 'Figure 4.6: Precision-Recall Curves for All Models')

    add_section_heading(doc, '4.7', 'Additional Evaluation Charts')

    add_image_if_exists(doc, 'calibration_curves.png', 5.0, 'Figure 4.7: Calibration Curves')
    add_image_if_exists(doc, 'learning_curves.png', 5.0, 'Figure 4.8: Learning Curves')
    add_image_if_exists(doc, 'lift_chart.png', 5.0, 'Figure 4.9: Lift Chart')
    add_image_if_exists(doc, 'cumulative_gain_chart.png', 5.0, 'Figure 4.10: Cumulative Gain Chart')
    add_image_if_exists(doc, 'feature_importance_comparison.png', 5.0, 'Figure 4.11: Feature Importance Comparison')

    add_body_text(doc, 'The lift chart shows that all models achieve lift values exceeding 2.0× in the top decile, meaning the highest-risk 10% of flights identified by the model contains more than double the proportion of actual delays compared to random selection. The Neural Network achieves the peak lift of approximately 2.48×, while XGBoost closely follows at 2.40×.')

    add_page_break(doc)

    # ============================
    # CHAPTER 5: CONCLUSION AND FUTURE WORK
    # ============================
    add_chapter_heading(doc, '5', 'CONCLUSION AND FUTURE WORK')

    add_section_heading(doc, '5.1', 'Conclusion')

    add_body_text(doc, 'This project presented AeroPredict, an explainable machine learning framework for flight delay prediction that combines high-accuracy classification with transparent, SHAP-based feature attribution. A comprehensive pipeline was developed encompassing a 25-dimensional feature space spanning meteorological, temporal, route-based, and operational variables, followed by median imputation, Random Forest–based feature selection, and standardized scaling.')

    add_body_text(doc, 'Six diverse models were trained and rigorously evaluated on a dataset of over 100,000 flight records: XGBoost, LightGBM, Random Forest, Deep Neural Network, Soft Voting Ensemble, and Stacking Ensemble. The key findings are summarized below:')

    conclusions = [
        'All six models surpassed 90% accuracy, with the Neural Network achieving the highest performance at 98.51% accuracy and 0.9883 AUC-ROC, validating the effectiveness of the proposed feature engineering pipeline.',
        'XGBoost (95.56% accuracy, 0.9842 AUC-ROC) was selected as the deployment model due to its optimal balance of accuracy, calibration, inference speed, and native compatibility with SHAP TreeExplainer.',
        'Five-fold stratified cross-validation confirmed model generalizability with standard deviations below ±0.005 across all models, indicating robust performance on unseen data.',
        'SHAP analysis consistently identified meteorological variables — precipitation, wind speed, and visibility — as the dominant delay predictors, aligning with domain knowledge and validating model trustworthiness.',
        'The AeroPredict AI web application was successfully deployed as a Flask-based interactive system supporting 70+ airports, automatic parameter computation, real-time weather-aware prediction, and per-prediction SHAP explanations.'
    ]
    for i, c in enumerate(conclusions, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Cm(1.27)
        run = p.add_run(f'{i}. {c}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    add_body_text(doc, 'The comprehensive evaluation through confusion matrices, ROC curves, precision-recall curves, calibration analysis, learning curves, lift charts, and radar charts confirmed the robustness, reliability, and calibration of the proposed framework. AeroPredict demonstrates that high-accuracy delay prediction and model interpretability are not mutually exclusive, and that the integration of XAI techniques into deployable systems is both feasible and valuable for safety-critical aviation applications.')

    add_section_heading(doc, '5.2', 'Future Work')

    add_body_text(doc, 'While the current framework demonstrates strong performance, several directions for future enhancement are identified:')

    future_items = [
        'Real-World Dataset Integration: Incorporating real flight records from the Bureau of Transportation Statistics (BTS), Flightradar24, and Indian DGCA databases to validate performance on authentic operational data.',
        'Real-Time Data APIs: Integration with live weather APIs (OpenWeatherMap, METAR), real-time flight tracking services (ADS-B), and airport operations databases for truly dynamic, real-time delay predictions.',
        'Temporal Deep Learning Models: Exploration of sequence-aware architectures such as LSTM, GRU, and Transformer-based models to capture delay propagation patterns across connecting flights.',
        'Attention-Based Explainability: Incorporating self-attention mechanisms within the neural network architecture to provide built-in interpretability as an alternative to post-hoc SHAP analysis.',
        'Multi-Class Delay Prediction: Extending the binary classification to a multi-class framework predicting delay severity categories — minor (15–30 min), moderate (30–60 min), significant (60–120 min), and severe (>120 min).',
        'Federated Learning: Implementing federated learning to enable multiple airlines to collaboratively train a shared model without exposing proprietary operational data.',
        'Mobile and Cloud Deployment: Extending to a mobile-responsive progressive web app (PWA) and deploying on cloud platforms (AWS/GCP) with auto-scaling.',
        'Causal Inference Integration: Moving beyond correlational feature importance to causal delay analysis using techniques such as DoWhy and causal forests.'
    ]

    for i, item in enumerate(future_items, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Cm(1.27)
        run = p.add_run(f'{i}. {item}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    add_page_break(doc)

    # ============================
    # CHAPTER 6: REFERENCES
    # ============================
    add_chapter_heading(doc, '6', 'REFERENCES')

    references = [
        '[1] International Air Transport Association (IATA), "Annual Review 2023: Air Passenger Market Analysis," IATA Economics Report, Montreal, Canada, 2023.',
        '[2] Federal Aviation Administration (FAA), "Cost of Delay Estimates: National Air Traffic Management System," Bureau of Transportation Statistics, U.S. Department of Transportation, Washington, D.C., 2022.',
        '[3] Eurocontrol, "CODA Digest: All-Causes Delay and Cancellations to Air Transport in Europe," Central Office for Delay Analysis Annual Report, Brussels, Belgium, 2022.',
        '[4] M. Ball, C. Barnhart, M. Dresner, M. Hansen, K. Neels, A. Odoni, E. Peterson, L. Sherry, A. Trani, and B. Zou, "Total Delay Impact Study: A Comprehensive Assessment of the Costs and Impacts of Flight Delay in the United States," NEXTOR Research Report, Univ. of Maryland, 2010.',
        '[5] R. Henriques and I. Feiteira, "Predictive Modelling: Flight Delays and Associated Factors, Hartsfield–Jackson Atlanta International Airport," Procedia Computer Science, vol. 136, pp. 214–221, 2018.',
        '[6] S. Choi, Y. J. Kim, S. Briceno, and D. Mavris, "Prediction of Weather-Induced Airline Delays Based on Machine Learning Algorithms," in Proc. IEEE/AIAA 35th Digital Avionics Systems Conference (DASC), Sacramento, CA, 2016, pp. 1–6.',
        '[7] J. J. Rebollo and H. Balakrishnan, "Characterization and Prediction of Air Traffic Delays," Transportation Research Part C: Emerging Technologies, vol. 44, pp. 231–241, 2014.',
        '[8] Y. J. Kim, S. Choi, S. Briceno, and D. Mavris, "A Deep Learning Approach for Flight Delay Prediction," in Proc. IEEE/AIAA 35th Digital Avionics Systems Conference (DASC), Sacramento, CA, 2016, pp. 1–8.',
        '[9] L. Breiman, "Random Forests," Machine Learning, vol. 45, no. 1, pp. 5–32, 2001.',
        '[10] T. Chen and C. Guestrin, "XGBoost: A Scalable Tree Boosting System," in Proc. 22nd ACM SIGKDD Int. Conf. on Knowledge Discovery and Data Mining (KDD), San Francisco, CA, 2016, pp. 785–794.',
        '[11] Y. LeCun, Y. Bengio, and G. Hinton, "Deep Learning," Nature, vol. 521, no. 7553, pp. 436–444, May 2015.',
        '[12] G. Ke, Q. Meng, T. Finley, T. Wang, W. Chen, W. Ma, Q. Ye, and T.-Y. Liu, "LightGBM: A Highly Efficient Gradient Boosting Decision Tree," in Advances in Neural Information Processing Systems (NeurIPS), vol. 30, 2017, pp. 3146–3154.',
        '[13] B. Ye, B. Liu, Y. Tian, and L. Wan, "A Methodology for Predicting Aggregate Flight Departure Delays in Airports Based on Supervised Learning," Sustainability, vol. 12, no. 7, p. 2749, 2020.',
        '[14] W. Jiang and J. Zhang, "Airline Delay Prediction Using Stacking Ensemble Learning," Journal of Air Transport Management, vol. 95, p. 102085, 2021.',
        '[15] C. Rudin, "Stop Explaining Black Box Machine Learning Models for High Stakes Decisions and Use Interpretable Models Instead," Nature Machine Intelligence, vol. 1, no. 5, pp. 206–215, 2019.',
        '[16] S. M. Lundberg and S.-I. Lee, "A Unified Approach to Interpreting Model Predictions," in Advances in Neural Information Processing Systems (NeurIPS), vol. 30, 2017, pp. 4765–4774.',
        '[17] M. T. Ribeiro, S. Singh, and C. Guestrin, "\'Why Should I Trust You?\': Explaining the Predictions of Any Classifier," in Proc. 22nd ACM SIGKDD Int. Conf. on Knowledge Discovery and Data Mining (KDD), San Francisco, CA, 2016, pp. 1135–1144.',
        '[18] S. Fernandes, M. Carvalho, and A. P. Teixeira, "Explainable Artificial Intelligence for Flight Trajectory Prediction," IEEE Trans. on Intelligent Transportation Systems, vol. 23, no. 10, pp. 19363–19374, 2022.',
        '[19] H. Wang, Z. Li, and Y. Zhang, "Interpretable Flight Delay Prediction with SHAP-Based Feature Attribution Analysis," Transportation Research Record, vol. 2677, no. 4, pp. 412–425, 2023.',
        '[20] S. Manna, S. Biswas, R. Kundu, S. Rakshit, P. Gupta, and S. Barman, "A Statistical Approach to Predict Flight Delay Using Gradient Boosted Decision Tree," in Proc. Int. Conf. on Computational Intelligence in Data Science (ICCIDS), Chennai, India, 2017, pp. 1–5.',
        '[21] N. Chakrabarty, "A Data Mining Approach to Flight Arrival Delay Prediction for American Airlines," in Proc. 9th Annual Information Technology, Electromechanical Engineering and Microelectronics Conference (IEMECON), Jaipur, India, 2019, pp. 102–107.',
        '[22] G. Gui, F. Liu, J. Sun, J. Yang, Z. Zhou, and D. Zhao, "Flight Delay Prediction Based on Aviation Big Data and Machine Learning," IEEE Trans. on Vehicular Technology, vol. 69, no. 1, pp. 140–150, Jan. 2020.',
        '[23] L. Hao, M. Hansen, Y. Zhang, and J. Post, "New York, New York: Two Ways of Estimating the Delay Impact of New York Airports," Transportation Research Part E: Logistics and Transportation Review, vol. 70, pp. 245–260, 2014.',
        '[24] D. P. Kingma and J. Ba, "Adam: A Method for Stochastic Optimization," in Proc. 3rd Int. Conf. on Learning Representations (ICLR), San Diego, CA, 2015.',
        '[25] F. Pedregosa et al., "Scikit-learn: Machine Learning in Python," Journal of Machine Learning Research, vol. 12, pp. 2825–2830, 2011.',
        '[26] C. Molnar, Interpretable Machine Learning: A Guide for Making Black Box Models Explainable, 2nd ed., 2022.',
        '[27] S. N. Srihari, "Ensemble Methods in Machine Learning," in Proc. Int. Workshop on Multiple Classifier Systems (MCS), Lecture Notes in Computer Science, vol. 1857, Springer, 2000, pp. 1–15.',
        '[28] M. Abadi et al., "TensorFlow: A System for Large-Scale Machine Learning," in Proc. 12th USENIX Symposium on Operating Systems Design and Implementation (OSDI), Savannah, GA, 2016, pp. 265–283.',
    ]

    for ref in references:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(ref)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

    add_page_break(doc)

    # ============================
    # APPENDIX A - SAMPLE CODE
    # ============================
    add_formatted_paragraph(doc, 'APPENDIX A – SAMPLE CODE', font_size=14, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=24, space_after=18)

    add_formatted_paragraph(doc, 'Key Code Snippets from the AeroPredict Framework', font_size=13, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6)

    add_formatted_paragraph(doc, 'A.1 XGBoost Model Training', font_size=12, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6)

    code1 = """import xgboost as xgb
from sklearn.model_selection import train_test_split, StratifiedKFold
from sklearn.preprocessing import StandardScaler
from sklearn.feature_selection import SelectFromModel
from sklearn.ensemble import RandomForestClassifier

# Feature Selection using Random Forest
selector = SelectFromModel(
    RandomForestClassifier(n_estimators=100, random_state=42),
    threshold='mean'
)
X_selected = selector.fit_transform(X_imputed, y)

# StandardScaler
scaler = StandardScaler()
X_scaled = scaler.fit_transform(X_selected)

# XGBoost Model
xgb_model = xgb.XGBClassifier(
    n_estimators=600, learning_rate=0.03, max_depth=7,
    subsample=0.8, colsample_bytree=0.8,
    reg_alpha=0.05, reg_lambda=1.0,
    scale_pos_weight=len(y[y==0])/len(y[y==1]),
    random_state=42, use_label_encoder=False,
    eval_metric='logloss'
)
xgb_model.fit(X_train, y_train)"""

    p = doc.add_paragraph()
    run = p.add_run(code1)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    add_formatted_paragraph(doc, 'A.2 SHAP Explainability Integration', font_size=12, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6)

    code2 = """import shap

# Initialize SHAP TreeExplainer
explainer = shap.TreeExplainer(xgb_model)

# Compute SHAP values for a prediction
shap_values = explainer.shap_values(X_scaled_instance)

# Extract top-5 contributing features
feature_importance = list(zip(feature_names, shap_values[0]))
feature_importance.sort(key=lambda x: abs(x[1]), reverse=True)
top_features = feature_importance[:5]

for name, value in top_features:
    direction = "increased" if value > 0 else "decreased"
    print(f"{name}: {direction} delay risk (SHAP: {value:.4f})")"""

    p = doc.add_paragraph()
    run = p.add_run(code2)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    add_formatted_paragraph(doc, 'A.3 Flask Prediction Endpoint', font_size=12, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6)

    code3 = """from flask import Flask, request, jsonify
import numpy as np

app = Flask(__name__)

@app.route('/predict', methods=['POST'])
def predict():
    data = request.json
    features = np.array([data['features']])
    features_scaled = scaler.transform(features)
    
    # Get prediction probability
    prob = xgb_model.predict_proba(features_scaled)[0][1]
    prediction = "DELAYED" if prob >= 0.5 else "ON TIME"
    
    # Risk level classification
    if prob < 0.3: risk = "LOW"
    elif prob < 0.5: risk = "MODERATE"
    elif prob < 0.8: risk = "HIGH"
    else: risk = "CRITICAL"
    
    # SHAP explanation
    shap_values = explainer.shap_values(features_scaled)
    
    return jsonify({
        'prediction': prediction,
        'probability': float(prob),
        'risk_level': risk,
        'shap_explanations': top_features
    })"""

    p = doc.add_paragraph()
    run = p.add_run(code3)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    add_page_break(doc)

    # ============================
    # APPENDIX B - OUTPUT SCREENSHOTS
    # ============================
    add_formatted_paragraph(doc, 'APPENDIX B – OUTPUT SCREENSHOTS', font_size=14, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=24, space_after=18)

    add_formatted_paragraph(doc, 'B.1 Model Training and Evaluation Output', font_size=12, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6)

    # Add all available chart images
    charts = [
        ('model_comparison.png', 'Figure B.1: Model Performance Comparison'),
        ('confusion_matrices.png', 'Figure B.2: Confusion Matrices'),
        ('roc_curves.png', 'Figure B.3: ROC Curves'),
        ('precision_recall_curves.png', 'Figure B.4: Precision-Recall Curves'),
        ('model_performance_radar.png', 'Figure B.5: Model Performance Radar Chart'),
        ('calibration_curves.png', 'Figure B.6: Calibration Curves'),
        ('learning_curves.png', 'Figure B.7: Learning Curves'),
        ('cross_validation_analysis.png', 'Figure B.8: Cross-Validation Analysis'),
        ('lift_chart.png', 'Figure B.9: Lift Chart'),
        ('cumulative_gain_chart.png', 'Figure B.10: Cumulative Gain Chart'),
        ('feature_importance_comparison.png', 'Figure B.11: Feature Importance Comparison'),
    ]

    for filename, caption in charts:
        add_image_if_exists(doc, filename, 5.0, caption)

    add_formatted_paragraph(doc, 'B.2 Web Application Screenshots', font_size=12, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6)

    # Try to add images from paper_images
    paper_imgs = [
        ('paper_images/image1.png', 'Figure B.12: AeroPredict AI Web Application Dashboard'),
        ('paper_images/image2.png', 'Figure B.13: Prediction Input Form'),
        ('paper_images/image3.png', 'Figure B.14: Prediction Results with SHAP Explanations'),
        ('paper_images/image4.png', 'Figure B.15: Research Dashboard'),
        ('paper_images/image5.png', 'Figure B.16: Model Evaluation Charts'),
    ]

    for filename, caption in paper_imgs:
        add_image_if_exists(doc, filename, 5.0, caption)

    add_page_break(doc)

    # ============================
    # APPENDIX C - OUTCOME ACHIEVED
    # ============================
    add_formatted_paragraph(doc, 'APPENDIX C – OUTCOME ACHIEVED', font_size=14, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=24, space_after=18)

    add_body_text(doc, 'The AeroPredict mini project successfully achieved the following outcomes:')

    outcomes = [
        'Development of a comprehensive machine learning pipeline for flight delay prediction incorporating 25 multi-dimensional features spanning meteorological, temporal, route, and operational categories.',
        'Training and rigorous evaluation of six distinct ML models (XGBoost, LightGBM, Random Forest, DNN, Soft Voting Ensemble, and Stacking Ensemble), all achieving above 90% accuracy with the Neural Network reaching 98.51%.',
        'Integration of SHAP (SHapley Additive exPlanations) for model interpretability, providing transparent per-prediction feature attribution.',
        'Deployment of a fully functional interactive web application (AeroPredict AI) using Flask, featuring real-time delay prediction, automatic parameter computation, weather integration, and SHAP-driven visual explanations.',
        'Support for 70+ airports worldwide (India, US, and International) with automatic Haversine distance computation and climate-zone-aware weather estimation.',
        'Comprehensive model validation through confusion matrices, ROC curves, precision-recall curves, calibration analysis, learning curves, lift charts, 5-fold cross-validation, and radar charts.',
        'Successful demonstration that high-accuracy predictive modeling and explainable AI are complementary rather than competing objectives in safety-critical aviation applications.'
    ]

    for i, outcome in enumerate(outcomes, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Cm(1.27)
        run = p.add_run(f'{i}. {outcome}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    add_formatted_paragraph(doc, 'Course Outcomes Mapped:', font_size=12, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=18, space_after=6)

    course_outcomes = [
        'CO1: Applied machine learning algorithms (XGBoost, LightGBM, Random Forest, DNN) to solve a real-world classification problem in the aviation domain.',
        'CO2: Implemented data preprocessing techniques including imputation, feature selection, and scaling for handling large-scale heterogeneous datasets.',
        'CO3: Evaluated and compared multiple model architectures using standardized metrics (Accuracy, Precision, Recall, F1-Score, AUC-ROC).',
        'CO4: Integrated Explainable AI (SHAP) to ensure model transparency and accountability in safety-critical decision-making.',
        'CO5: Designed and deployed a full-stack web application bridging the gap between machine learning research and practical user-facing tools.',
    ]

    for co in course_outcomes:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Cm(1.27)
        run = p.add_run(co)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    # ============================
    # SAVE THE DOCUMENT
    # ============================
    output_path = os.path.join(BASE_DIR, 'AeroPredict_Mini_Project_Report.docx')
    doc.save(output_path)
    print(f'Report generated successfully: {output_path}')
    return output_path


if __name__ == '__main__':
    generate_report()
